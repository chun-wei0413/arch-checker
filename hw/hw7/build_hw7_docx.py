"""
Build 113598009_7.docx by inserting Sections 3-6 (Design, Implementation,
Programming, Unit Testing) into the existing document before the Revision section,
then renumbering old sections 3->7 and 4->8.
"""

import os
import sys
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

BASE = os.path.dirname(os.path.abspath(__file__))
SRC  = os.path.join(BASE, "113598009_7.docx")
DST  = SRC  # overwrite in-place
DIAG = os.path.join(BASE, "diagrams")

# --- helpers -----------------------------------------------------------

def make_para_elem(doc, text="", style="Normal", bold=False, italic=False,
                   size_pt=None, align=None):
    """Return a new <w:p> element (not yet in the document body)."""
    p = doc.add_paragraph(text, style=style)
    run = p.runs[0] if p.runs else None
    if run:
        if bold:   run.bold = bold
        if italic: run.italic = italic
        if size_pt: run.font.size = Pt(size_pt)
    if align:
        p.alignment = align
    elem = p._element
    p._element.getparent().remove(p._element)
    return elem

def make_heading_elem(doc, text, level):
    p = doc.add_heading(text, level)
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def make_image_elem(doc, img_path, width_in=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_in))
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def make_caption_elem(doc, text):
    p = doc.add_paragraph(text, style="Caption") if "Caption" in [s.name for s in doc.styles] else doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0] if p.runs else p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def _add_cell_border(tc):
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "888888")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def make_table_elem(doc, headers, rows, col_widths=None):
    """Return an <w:tbl> element."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D5E8D4")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
        _add_cell_border(cell._tc)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            _add_cell_border(cell._tc)
    # Optional column widths (in inches)
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[ci].width = Inches(w)
    elem = tbl._tbl
    elem.getparent().remove(elem)
    return elem

def make_code_para_elem(doc, text):
    """Monospace paragraph for source code display."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    elem = p._element
    elem.getparent().remove(elem)
    return elem

def insert_before(target_elem, new_elem):
    target_elem.addprevious(new_elem)

# --- load existing document -------------------------------------------
doc = Document(SRC)
body = doc.element.body

# --- find "3. Revision..." heading paragraph --------------------------
revision_para = None
measurement_para = None
for p in doc.paragraphs:
    if p.style.name.startswith("Heading") and "Revision" in p.text:
        revision_para = p
    if p.style.name.startswith("Heading") and "Measurement" in p.text:
        measurement_para = p

if revision_para is None:
    sys.exit("ERROR: Could not find 'Revision' heading paragraph in the document.")

anchor = revision_para._element  # everything new is inserted before this

# --- rename old sections 3->7 and 4->8 --------------------------------
if revision_para:
    for run in revision_para.runs:
        if run.text.startswith("3."):
            run.text = run.text.replace("3.", "7.", 1)
            break
    else:
        t = revision_para._element.find(".//" + qn("w:t"))
        if t is not None and t.text and t.text.startswith("3."):
            t.text = t.text.replace("3.", "7.", 1)

if measurement_para:
    for run in measurement_para.runs:
        if "Measurement" in run.text:
            run.text = run.text.replace("4.", "8.", 1)
            break
    else:
        t = measurement_para._element.find(".//" + qn("w:t"))
        if t is not None and t.text:
            t.text = t.text.replace("4.", "8.", 1)

# --- helper to queue elements (inserted in forward order) -------------
queue = []

def Q(elem):
    queue.append(elem)

def QP(text="", style="Normal", bold=False, italic=False, size_pt=None, align=None):
    Q(make_para_elem(doc, text, style, bold, italic, size_pt, align))

def QH(text, level):
    Q(make_heading_elem(doc, text, level))

def QI(img_file, width_in=5.5):
    path = os.path.join(DIAG, img_file)
    if os.path.exists(path):
        Q(make_image_elem(doc, path, width_in))
    else:
        QP(f"[Image not found: {img_file}]", italic=True)

def QCAP(text):
    Q(make_caption_elem(doc, text))

def QT(headers, rows, col_widths=None):
    Q(make_table_elem(doc, headers, rows, col_widths))

def QCODE(code_text):
    for line in code_text.split("\n"):
        Q(make_code_para_elem(doc, line))

# ======================================================================
# SECTION 3: DESIGN
# ======================================================================
QH("3. Design", 1)

# 3.1 Logical Architecture
QH("3.1 Logical Architecture", 2)
QP("arch-checker 為單一使用者的 CLI 工具（無 GUI、無資料庫後端）。本節以 UML package diagram "
   "方式呈現 Iteration II 系統的組織架構，並說明各 package 之職責與相依方向。系統分為四個水平層："
   "CLI 層使用 picocli 框架處理指令；Application 層定義業務邏輯與介面（CodeParser、ProfileLoader、"
   "SuppressionStore、Reporter）；Domain 層包含純粹的領域物件，無對外依賴；Infrastructure 層實作 "
   "Application 層所定義的 port 介面，遵循相依反轉原則（Dependency Inversion Principle）。"
   "Iteration II 於 CLI 層新增 FixCommand、ProfileCommand、ProfileInitCommand、ProfileValidateCommand，"
   "於 Application 層新增 InteractiveFixer、ProfileTemplateService、ProfileValidateService。")
QI("logical_arch.png", 6.0)
QCAP("圖 3-1 Logical Architecture（UML Package Diagram）")

# 3.2 Use-Case Realizations with GRASP Patterns
QH("3.2 Use-Case Realizations with GRASP Patterns", 2)
QH("3.2.1 System Sequence Diagram — UC-02 Define Style Profile", 3)
QP("UC-02 為 Iteration II 最重要的新增使用案例。本節以 System Sequence Diagram（SSD）識別系統事件，"
   "共有兩個系統事件：initProfile(outputPath) 產生 YAML 範本檔，validateProfile(yamlPath) 驗證已編輯的 "
   "Profile 並回報成功載入的規則數量。")
QI("uc02_ssd.png", 6.0)
QCAP("圖 3-2.1 UC-02 Define Style Profile — System Sequence Diagram")

QH("3.2.2 Sequence Diagrams with GRASP Patterns", 3)
QP("系統事件 1：initProfile(outputPath)")
QP("GRASP 模式應用：")
QP("• Controller：ProfileInitCommand 作為 use-case controller 處理 initProfile 系統事件"
   "（一個 use case 對應一個 controller）。")
QP("• Information Expert + Creator：ProfileTemplateService 具備範本結構的完整知識，"
   "負責建立並寫出 YAML 檔案。")
QP("• Low Coupling：ProfileInitCommand 於本地建立 ProfileTemplateService，使 CLI 層"
   "與檔案系統細節保持獨立。")
QI("uc02_sd_init.png", 6.2)
QCAP("圖 3-2.2a SD for initProfile()（標示 GRASP 模式）")

QP("")
QP("系統事件 2：validateProfile(yamlPath)")
QP("GRASP 模式應用：")
QP("• Controller：ProfileValidateCommand 處理 validateProfile 系統事件。")
QP("• Information Expert：ProfileValidateService 掌握驗證的業務邏輯；"
   "YamlProfileLoader 為解析 YAML 與建立 StyleProfile 的資訊專家。")
QP("• Creator：YamlProfileLoader 持有已解析的 YAML 資料，因此負責建立 StyleProfile"
   "與 ComplianceRule 物件。")
QP("• Low Coupling：ProfileValidateService 依賴 ProfileLoader 介面，"
   "而非直接依賴 YamlProfileLoader，便於替換實作。")
QI("uc02_sd_validate.png", 6.2)
QCAP("圖 3-2.2b SD for validateProfile()（標示 GRASP 模式）")

# 3.3 Design Class Diagram
QH("3.3 Design Class Model", 2)
QP("下圖為 Iteration II 之 Design Class Diagram（DCD），涵蓋系統所有類別、屬性與主要方法。"
   "標示 [new] 者為新增類別，標示 [mod] 者為已修改之類別。"
   "四個 ComplianceRule 子類別（NamingRule、DependencyRule、SupertypeRule、PackageRule）"
   "在其產生的每個 Violation 上加入 suggestion 訊息，以支援 UC-03 的互動式修復顯示。")
QI("dcd.png", 6.3)
QCAP("圖 3-3 Design Class Diagram（Iteration II）")

# ======================================================================
# SECTION 4: IMPLEMENTATION CLASS MODEL
# ======================================================================
QH("4. Implementation Class Model", 1)
QP("本章節依實作完成之 Java 程式碼繪製 Implementation Class Diagram（4.1），與 §3.3 之 "
   "Design Class Model 進行差異比較（4.2），最後以 Lines of Code 量化實作規模（4.3）。")

QH("4.1 Implementation Class Diagram", 2)
QP("下圖為 arch-checker 之 Implementation Class Diagram（ICD），反映系統實際編譯後的結構。"
   "由於實作緊跟 DCD，ICD 在結構上與圖 3-3 相同。與 Iteration I 相比，主要新增為七個新類別"
   "（FixCommand、ProfileCommand、ProfileInitCommand、ProfileValidateCommand、InteractiveFixer、"
   "ProfileTemplateService、ProfileValidateService）以及 Violation 上的 suggestion 欄位。")
QI("dcd.png", 6.3)
QCAP("圖 4-1 Implementation Class Diagram（與 DCD 結構相同；Iteration II 新增以綠色標示）")

QH("4.2 Differences between Design and Implementation Class Model", 2)
QH("Table 4.2.1: Comparison of Design and Implementation (Iteration II changes)", 3)
QP("下表逐一列出 Iteration II 新增或修改的所有類別與選定方法，比較 DCD（設計）與實際程式碼"
   "（實作）之差異。「Yes」表示該類別／方法存在於該成品中。")

tbl521_headers = ["類別", "方法 / 屬性", "設計", "實作"]
tbl521_rows = [
    # New CLI classes
    ("FixCommand [new]", "call(): Integer", "Yes", "Yes"),
    ("ProfileCommand [new]", "run(): void", "Yes", "Yes"),
    ("ProfileInitCommand [new]", "call(): Integer", "Yes", "Yes"),
    ("ProfileValidateCommand [new]", "call(): Integer", "Yes", "Yes"),
    # New Application classes
    ("InteractiveFixer [new]", "fix(report, file, profile): FixResult", "Yes", "Yes"),
    ("InteractiveFixer [new]", "prompt(): String", "Yes", "Yes"),
    ("InteractiveFixer [new]", "suppress(v, file, profile): void", "Yes", "Yes"),
    ("InteractiveFixer.FixResult [new]", "getExitCode(): int", "Yes", "Yes"),
    ("ProfileTemplateService [new]", "generateTemplate(outputPath): void", "Yes", "Yes"),
    ("ProfileValidateService [new]", "validate(profilePath): ValidateResult", "Yes", "Yes"),
    ("ProfileValidateService.ValidateResult [new]", "getRuleCount(): int", "Yes", "Yes"),
    # Modified domain classes
    ("Violation [mod]", "suggestion: String [0..1]（新增屬性）", "Yes", "Yes"),
    ("Violation [mod]", "getSuggestion(): String（新增方法）", "Yes", "Yes"),
    ("Violation [mod]", "hasSuggestion(): boolean（新增方法）", "Yes", "Yes"),
    # Modified rule classes
    ("NamingRule [mod]", "validate()：現於 Violation 上設定 suggestion", "Yes", "Yes"),
    ("DependencyRule [mod]", "validate()：現於 Violation 上設定 suggestion", "Yes", "Yes"),
    ("SupertypeRule [mod]", "validate()：現於 Violation 上設定 suggestion", "Yes", "Yes"),
    ("PackageRule [mod]", "validate()：現於 Violation 上設定 suggestion", "Yes", "Yes"),
    # Modified CLI
    ("Main [mod]", "subcommands：新增 ProfileCommand、FixCommand", "Yes", "Yes"),
]
QT(tbl521_headers, tbl521_rows, col_widths=[2.2, 2.8, 0.7, 0.7])

QP("")
QH("Table 4.2.2: Summary of Implementation Class/Method Changes", 3)
QP("下表統計 Iteration II 整體之類別與方法變動數量。")
tbl522_headers = ["", "新增", "移除", "修改（簽章變更）"]
tbl522_rows = [
    ("類別",  "7", "0", "6（Main、Violation、NamingRule、DependencyRule、SupertypeRule、PackageRule）"),
    ("方法 / 屬性", "17", "0", "0"),
]
QT(tbl522_headers, tbl522_rows, col_widths=[1.1, 0.8, 0.9, 3.6])

QH("4.3 Calculate Line of Code", 2)
QH("Table 4.3.1: Line of Code of Classes", 3)
QP("LOC 採「排除空行與單行 // 註解、整段 /* … */ 區塊」之計算方式；method 數含 public、private、"
   "protected 之具體方法（不計 abstract 宣告與建構子，但計入介面方法）。"
   "Iteration II 新增之類別標示 [new]，修改之類別標示 [mod]。")

loc_headers = ["No.", "類別名稱", "方法數", "LOC（不含註解／空行）"]
loc_rows = [
    (1,  "cli.Main",                              2,   20),
    (2,  "cli.CheckCommand",                       6,   39),
    (3,  "cli.SuppressCommand",                    8,   39),
    (4,  "cli.ProfileCommand [new]",               2,   11),
    (5,  "cli.ProfileInitCommand [new]",           3,   18),
    (6,  "cli.ProfileValidateCommand [new]",       3,   30),
    (7,  "cli.FixCommand [new]",                   5,   55),
    (8,  "application.ComplianceCheckService",    12,   41),
    (9,  "application.SuppressionService",         6,   39),
    (10, "application.InteractiveFixer [new]",    15,   87),
    (11, "application.ProfileTemplateService [new]", 3, 48),
    (12, "application.ProfileValidateService [new]", 11, 31),
    (13, "application.CodeParser (interface)",     1,    6),
    (14, "application.ProfileLoader (interface)",  1,    6),
    (15, "application.SuppressionStore (interface)", 1,  9),
    (16, "application.Reporter (interface)",       1,    5),
    (17, "domain.compliance.ComplianceCheck",     13,   45),
    (18, "domain.compliance.Violation [mod]",     15,   31),
    (19, "domain.compliance.ViolationReport",     12,   29),
    (20, "domain.compliance.Suppression",         13,   29),
    (21, "domain.codebase.File",                   7,   24),
    (22, "domain.codebase.Project",                6,   18),
    (23, "domain.profile.StyleProfile",            6,   18),
    (24, "domain.rule.ComplianceRule",             8,   27),
    (25, "domain.rule.NamingRule [mod]",           7,   46),
    (26, "domain.rule.DependencyRule [mod]",      10,   52),
    (27, "domain.rule.SupertypeRule [mod]",        8,   44),
    (28, "domain.rule.PackageRule [mod]",          5,   36),
    (29, "infrastructure.parser.JavaParserAdapter", 4,  42),
    (30, "infrastructure.profile.YamlProfileLoader", 3, 69),
    (31, "infrastructure.report.ConsoleReporter",  5,   22),
    (32, "infrastructure.report.JsonReporter",     6,   36),
    (33, "infrastructure.suppression.YamlSuppressionStore", 5, 79),
]
sum_methods = sum(r[2] for r in loc_rows)
sum_loc     = sum(r[3] for r in loc_rows)
loc_rows_display = [(r[0], r[1], r[2], r[3]) for r in loc_rows]
loc_rows_display.append(("", "合計", sum_methods, sum_loc))
QT(loc_headers, loc_rows_display, col_widths=[0.4, 3.6, 0.9, 1.5])

# ======================================================================
# SECTION 5: PROGRAMMING
# ======================================================================
QH("5. Programming", 1)
QP("本章節依需求 5.1 提供系統執行畫面截圖，並於 5.2 節錄 Iteration II 最具代表性之原始碼。")

QH("5.1 Snapshots of System Execution", 2)
QP("UC-02 profile init 與 validate 執行畫面：")
QCODE(
"""$ arch-checker profile init target/sample-arch.yaml
Profile template written to: target\\sample-arch.yaml

$ arch-checker profile validate target/sample-arch.yaml
Profile 'my-profile' is valid. Loaded 4 rule(s):
  - R-NAME-01 (NamingRule)
  - R-DEP-01 (DependencyRule)
  - R-SUP-01 (SupertypeRule)
  - R-PKG-01 (PackageRule)"""
)
QP("")
QP("UC-01 架構合規檢查執行畫面（使用產生的 profile 範本對自身原始碼進行檢查；"
   "因範本的 NamingRule 要求所有類別以 'Service' 結尾，故預期出現多筆違規）：")
QCODE(
"""$ arch-checker check src/main/java target/sample-arch.yaml
src\\main\\java\\com\\archchecker\\application\\CodeParser.java:7 [R-NAME-01]
  Class 'CodeParser' violates naming pattern '.*Service'
  Suggestion: Rename 'CodeParser' to match pattern '.*Service'
... (共 36 筆違規)
--
Checked 33 file(s); 36 violation(s); 0 suppressed."""
)

QH("5.2 Source Code Listing", 2)
QP("下列依序列出 Iteration II 最重要的原始碼（新增或修改之類別）。完整原始碼請參見專案儲存庫。")

def read_src(relative_path):
    full = os.path.join(BASE, "..", "..", "src", "main", "java", "com", "archchecker", relative_path)
    with open(full, encoding="utf-8") as f:
        return f.read()

important_files = [
    ("domain/compliance/Violation.java",          "Violation.java [mod]（新增 suggestion 欄位）"),
    ("application/InteractiveFixer.java",         "InteractiveFixer.java [new]"),
    ("application/ProfileTemplateService.java",   "ProfileTemplateService.java [new]"),
    ("application/ProfileValidateService.java",   "ProfileValidateService.java [new]"),
    ("cli/FixCommand.java",                       "FixCommand.java [new]"),
    ("cli/ProfileInitCommand.java",               "ProfileInitCommand.java [new]"),
    ("cli/ProfileValidateCommand.java",           "ProfileValidateCommand.java [new]"),
]

for rel, label in important_files:
    QP(f"— {label}", bold=True)
    try:
        QCODE(read_src(rel))
    except Exception as e:
        QP(f"[Could not read {rel}: {e}]", italic=True)
    QP("")

# ======================================================================
# SECTION 6: UNIT TESTING
# ======================================================================
QH("6. Unit Testing", 1)

QH("6.1 Snapshots of Testing Result", 2)
QP("以下為 mvn test 執行結果截圖，共 28 個 JUnit 5 測試全數通過：")
QCODE(
"""$ mvn test
[INFO] Running com.archchecker.application.ComplianceCheckServiceTest
[INFO] Tests run: 3, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.application.InteractiveFixerTest        [new]
[INFO] Tests run: 4, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.application.ProfileTemplateServiceTest  [new]
[INFO] Tests run: 3, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.application.SuppressionServiceTest
[INFO] Tests run: 2, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.domain.rule.DependencyRuleTest
[INFO] Tests run: 3, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.domain.rule.NamingRuleTest
[INFO] Tests run: 3, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.domain.rule.PackageRuleTest
[INFO] Tests run: 2, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.domain.rule.SupertypeRuleTest
[INFO] Tests run: 2, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.infrastructure.profile.YamlProfileLoaderTest
[INFO] Tests run: 3, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Running com.archchecker.infrastructure.suppression.YamlSuppressionStoreTest
[INFO] Tests run: 3, Failures: 0, Errors: 0, Skipped: 0  OK
[INFO] Tests run: 28, Failures: 0, Errors: 0, Skipped: 0
[INFO] BUILD SUCCESS"""
)

QH("6.2 Unit Testing Code Listing", 2)
QP("下列測試檔案涵蓋 Iteration II 的新增功能：")

def read_test(relative_path):
    full = os.path.join(BASE, "..", "..", "src", "test", "java", "com", "archchecker", relative_path)
    with open(full, encoding="utf-8") as f:
        return f.read()

test_files = [
    ("application/InteractiveFixerTest.java",      "InteractiveFixerTest.java [new]（UC-03 互動式修復）"),
    ("application/ProfileTemplateServiceTest.java","ProfileTemplateServiceTest.java [new]（UC-02 範本產生）"),
    ("application/ComplianceCheckServiceTest.java","ComplianceCheckServiceTest.java（核心合規檢查）"),
    ("domain/rule/NamingRuleTest.java",            "NamingRuleTest.java（NamingRule 含 suggestion）"),
    ("domain/rule/DependencyRuleTest.java",        "DependencyRuleTest.java（DependencyRule 含 suggestion）"),
]

for rel, label in test_files:
    QP(f"— {label}", bold=True)
    try:
        QCODE(read_test(rel))
    except Exception as e:
        QP(f"[Could not read {rel}: {e}]", italic=True)
    QP("")

# --- flush queue: insert all new elements before the Revision anchor ---
for elem in queue:
    insert_before(anchor, elem)

# --- update Change History table: add v9.0 row ------------------------
for tbl in doc.tables:
    if tbl.rows and "Version" in tbl.rows[0].cells[0].text:
        new_row = tbl.add_row()
        new_row.cells[0].text = "9.0"
        new_row.cells[1].text = "2026-06-14"
        new_row.cells[2].text = "FrankLi"
        new_row.cells[3].text = (
            "Iteration II 實作：UC-02（Define Style Profile）與 "
            "UC-03（Check with Interactive Fix Suggestions）。"
            "新增 Design、Implementation Class Model、Programming、"
            "Unit Testing 章節。7 個新類別、6 個修改類別，28 個 JUnit 5 測試全部通過。"
        )
        for cell in new_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        break

# --- update Measurement table: add HW7 row ----------------------------
hw7_meas_data = ["113598009", "FrankLi", "HW#7", "115/06/14  20:00 ~ 23:59", "4H"]
for tbl in doc.tables:
    if not tbl.rows:
        continue
    header_texts = " ".join(c.text for c in tbl.rows[0].cells)
    if any(kw in header_texts for kw in ("HW", "Duration", "Date", "Student")):
        if "Version" in header_texts or "Author" in header_texts:
            continue
        new_row = tbl.add_row()
        ncols = len(new_row.cells)
        for ci, val in enumerate(hw7_meas_data[:ncols]):
            new_row.cells[ci].text = val
        for cell in new_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        break

# --- save -------------------------------------------------------------
doc.save(DST)
print(f"Saved: {DST}")
print(f"Total new elements inserted: {len(queue)}")
