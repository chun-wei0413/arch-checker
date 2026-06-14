"""
Remove all content inserted between Domain Model and Revision sections (indices 78-963).
Leaves the original attributes table at [77] and Revision section at [964+] intact.
"""

import os
from docx import Document
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
SRC  = os.path.join(BASE, "113598009_7.docx")

doc  = Document(SRC)
body = doc.element.body
children = list(body)

# Find Revision heading (H1)
revision_idx = None
for i, ch in enumerate(children):
    if ch.tag == qn("w:p"):
        pPr = ch.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None and pStyle.get(qn("w:val"), "") in ("1", "Heading1"):
                texts = ch.findall(".//" + qn("w:t"))
                txt = "".join(t.text or "" for t in texts)
                if "Revision" in txt:
                    revision_idx = i
                    print(f"Revision heading at [{i}]: {txt[:60]}")
                    break

# Delete from index 78 up to (not including) revision_idx
delete_start = 78
delete_end   = revision_idx  # exclusive

print(f"Deleting indices {delete_start} to {delete_end - 1} ({delete_end - delete_start} elements)")

to_delete = children[delete_start:delete_end]
for elem in to_delete:
    body.remove(elem)

doc.save(SRC)

# Verify
doc2 = Document(SRC)
children2 = list(doc2.element.body)
print(f"New body size: {len(children2)}")
headings = [p.text for p in doc2.paragraphs if p.style.name.startswith("Heading 1")]
print("H1 headings:")
for h in headings:
    print(f"  {repr(h)}")
